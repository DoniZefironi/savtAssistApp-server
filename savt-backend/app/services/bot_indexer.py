import asyncio
import logging
from pathlib import Path

from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.models.embedding import EMBEDDING_DIM, Embedding
from app.models.faq_entry import FaqEntry
from app.models.kbarticle import KbArticle
from app.models.kb_article_attachment import KbArticleAttachment
from app.services import yandex_service

logger = logging.getLogger(__name__)

UPLOAD_ROOT = Path("/code/uploads")
CHUNK_SIZE = 800
CHUNK_OVERLAP = 100



def _chunks(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    result = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        result.append(text[start:end])
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return result



async def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
    except Exception:
        logger.exception("Не удалось открыть PDF: %s", path)
        return ""

    parts: list[str] = []
    for page_num, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
            continue
        # Страница без текстового слоя — похоже на скан. Достаём встроенные
        # изображения страницы напрямую через pypdf (без poppler/pdf2image)
        # и распознаём их через Yandex Vision OCR.
        try:
            images = list(page.images)
        except Exception:
            images = []
        for img in images:
            try:
                ocr_text = await yandex_service.ocr_image(img.data)
                if ocr_text.strip():
                    parts.append(ocr_text)
            except Exception:
                logger.exception(
                    "OCR не удался для страницы %d файла %s", page_num + 1, path
                )
            await asyncio.sleep(0.2)
    return "\n\n".join(parts)


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs]
        # Технические характеристики в таких документах часто оформлены
        # таблицами — doc.paragraphs их не видит, обходим отдельно.
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception:
        logger.exception("Не удалось разобрать Word-документ: %s", path)
        return ""


def _parse_excel(path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        parts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        logger.exception("Не удалось разобрать Excel-файл: %s", path)
        return ""


async def _ocr_image_file(path: Path) -> str:
    try:
        return await yandex_service.ocr_image(path.read_bytes())
    except Exception:
        logger.exception("OCR не удался для изображения: %s", path)
        return ""


_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp"}

# Старые бинарные форматы (Word/Excel 97-2003) современными библиотеками не
# читаются в принципе — нужна конвертация в .docx/.xlsx, а не парсинг.
_UNSUPPORTED_LEGACY_FORMATS = {
    ".doc": "Word 97-2003 (.doc) — пересохраните файл как .docx",
    ".xls": "Excel 97-2003 (.xls) — пересохраните файл как .xlsx",
}


async def _extract_text(file_url: str) -> str:
    relative = file_url.removeprefix("/static/")
    path = UPLOAD_ROOT / relative
    if not path.exists():
        logger.warning("Файл для индексации не найден на диске: %s", path)
        return ""

    suffix = path.suffix.lower()
    if suffix in _UNSUPPORTED_LEGACY_FORMATS:
        logger.warning(
            "Формат не поддерживается извлечением текста (%s): %s",
            _UNSUPPORTED_LEGACY_FORMATS[suffix], path,
        )
        return ""

    if suffix == ".pdf":
        text = await _parse_pdf(path)
    elif suffix == ".docx":
        text = _parse_docx(path)
    elif suffix == ".xlsx":
        text = _parse_excel(path)
    elif suffix in _IMAGE_SUFFIXES:
        text = await _ocr_image_file(path)
    else:
        logger.info("Формат %s не поддерживается для извлечения текста: %s", suffix, path)
        return ""

    if not text.strip():
        logger.warning("Извлечённый текст пуст (файл без текстового слоя?): %s", path)
    return text



async def _upsert_chunks(
    session: AsyncSession,
    source_type: str,
    source_id: int,
    chunks: list[str],
    meta: dict,
) -> None:
    await session.execute(
        delete(Embedding).where(
            Embedding.source_type == source_type,
            Embedding.source_id == source_id,
        )
    )
    for i, chunk in enumerate(chunks):
        vector = await yandex_service.embed_document(chunk)
        session.add(Embedding(
            source_type=source_type,
            source_id=source_id,
            chunk_index=i,
            content=chunk,
            embedding=vector,
            meta=meta,
        ))
        # Yandex лимит: 10 запросов/сек → 0.12с между запросами = ~8/сек
        await asyncio.sleep(0.12)
    await session.flush()



async def index_faq_entry(session: AsyncSession, entry: FaqEntry) -> None:
    text = f"Вопрос: {entry.question}\nОтвет: {entry.answer}"
    await _upsert_chunks(session, "faq", entry.id, _chunks(text), {"title": entry.question})


async def index_kb_article(session: AsyncSession, article: KbArticle) -> None:
    parts = []
    if article.title:
        parts.append(article.title)
    if article.content:
        parts.append(article.content)

    attachments = (await session.execute(
        select(KbArticleAttachment).where(KbArticleAttachment.article_id == article.id)
    )).scalars().all()
    for att in attachments:
        parts.append(await _extract_text(att.file_url))

    text = "\n\n".join(p for p in parts if p.strip())
    await _upsert_chunks(session, "kb_article", article.id, _chunks(text), {"title": article.title})


async def index_document(session: AsyncSession, doc: Document) -> None:
    # Служебный документ (is_internal) не должен быть виден пользователю вообще —
    # значит, и в поиске бота его содержимое всплывать не должно. Чистим любые
    # уже существующие эмбеддинги (документ мог быть проиндексирован раньше, до
    # того как его закрыли) и не создаём новые — это же убирает из поиска файлы,
    # подхваченные с NAS напрямую (см. import_new_files_from_nas), у них
    # is_internal=True с самого создания.
    if doc.is_internal:
        await session.execute(
            delete(Embedding).where(
                Embedding.source_type == "document",
                Embedding.source_id == doc.id,
            )
        )
        return
    text = await _extract_text(doc.file_url)
    # Не удалось извлечь текст (нераспознанный формат, битый файл, документ без
    # текстового слоя и т.п.) — индексируем хотя бы заголовок, чтобы бот вообще
    # знал о существовании файла, но помечаем это в meta. Без этой пометки такой
    # документ навсегда застревал бы "уже проиндексированным" (у него ведь есть
    # embeddings) — reindex_all(force=False) больше никогда не пытался бы
    # переизвлечь текст повторно, даже после починки самого экстрактора.
    extraction_failed = not text.strip()
    if extraction_failed:
        text = doc.title or ""
    await _upsert_chunks(
        session, "document", doc.id, _chunks(text),
        {"title": doc.title, "cabinet_id": doc.cabinet_id, "extraction_failed": extraction_failed},
    )


def schedule_reindex_document(doc_id: int) -> None:
    """Фоновая (best-effort) переиндексация одного документа — после создания
    через форму загрузки или после смены is_internal через PATCH. Отдельная
    сессия — вызывается уже после commit основного запроса, не должна его
    блокировать/ронять."""
    async def _task():
        from app.database import AsyncSessionLocal
        try:
            async with AsyncSessionLocal() as s:
                doc = await s.get(Document, doc_id)
                if doc:
                    await index_document(s, doc)
                    await s.commit()
        except Exception:
            logger.exception("Фоновая переиндексация документа %s не удалась", doc_id)
    asyncio.create_task(_task())


async def reindex_all(session: AsyncSession, force: bool = False) -> dict:
    """Индексирует только ещё не проиндексированные записи.
    force=True — переиндексирует всё (старое поведение).

    Каждый элемент коммитится отдельно и сам ловит свою ошибку: раньше весь
    прогон коммитился одним разом в конце, и сбой на одном документе (например,
    сетевая ошибка Yandex API на одном из многих чанков большого файла) откатывал
    вообще всё, включая уже успешно проиндексированные до него FAQ/статьи/
    документы этого же прогона. Теперь одна неудача попадает в stats["failed"]
    и не мешает остальным."""
    stats = {"faq": 0, "kb_article": 0, "document": 0, "skipped": 0, "failed": 0}

    if not force:
        rows = (await session.execute(
            select(Embedding.source_type, Embedding.source_id, Embedding.meta)
        )).all()
        # Документы с extraction_failed=True в already НЕ попадают — иначе
        # обычный "переиндексировать новое" (force=False) навсегда пропускал бы
        # документ, у которого когда-то не извлёкся текст (см. index_document)
        already = {
            (source_type, source_id) for source_type, source_id, meta in rows
            if not (source_type == "document" and (meta or {}).get("extraction_failed"))
        }
    else:
        already = set()

    entries = (await session.execute(select(FaqEntry))).scalars().all()
    for e in entries:
        if ("faq", e.id) in already:
            stats["skipped"] += 1
            continue
        try:
            await index_faq_entry(session, e)
            await session.commit()
            stats["faq"] += 1
        except Exception:
            await session.rollback()
            stats["failed"] += 1
            logger.exception("Индексация FAQ %s не удалась", e.id)

    articles = (await session.execute(select(KbArticle))).scalars().all()
    for a in articles:
        if ("kb_article", a.id) in already:
            stats["skipped"] += 1
            continue
        try:
            await index_kb_article(session, a)
            await session.commit()
            stats["kb_article"] += 1
        except Exception:
            await session.rollback()
            stats["failed"] += 1
            logger.exception("Индексация статьи КБ %s не удалась", a.id)

    docs = (await session.execute(select(Document))).scalars().all()
    for d in docs:
        if ("document", d.id) in already:
            stats["skipped"] += 1
            continue
        try:
            await index_document(session, d)
            await session.commit()
            stats["document"] += 1
        except Exception:
            await session.rollback()
            stats["failed"] += 1
            logger.exception("Индексация документа %s не удалась", d.id)

    return stats


_MODEL_BY_SOURCE_TYPE = {
    "faq": FaqEntry,
    "kb_article": KbArticle,
    "document": Document,
}


async def prune_orphaned(session: AsyncSession) -> dict:
    """Удаляет embeddings, чей источник (FAQ/статья КБ/документ) больше не
    существует. Основной сценарий — удаление категории каскадно сносит её
    статьи/вопросы на уровне БД (ondelete=CASCADE), в обход сервисного
    delete(), который обычно чистит embeddings сам."""
    stats = {"faq": 0, "kb_article": 0, "document": 0}

    rows = (await session.execute(
        select(Embedding.source_type, Embedding.source_id).distinct()
    )).all()
    ids_by_type: dict[str, set[int]] = {}
    for source_type, source_id in rows:
        ids_by_type.setdefault(source_type, set()).add(source_id)

    for source_type, ids in ids_by_type.items():
        model = _MODEL_BY_SOURCE_TYPE.get(source_type)
        if model is None:
            continue
        existing_ids = set((await session.execute(
            select(model.id).where(model.id.in_(ids))
        )).scalars().all())
        orphan_ids = ids - existing_ids
        if not orphan_ids:
            continue
        await session.execute(
            delete(Embedding).where(
                Embedding.source_type == source_type,
                Embedding.source_id.in_(orphan_ids),
            )
        )
        stats[source_type] = len(orphan_ids)
        logger.info("Удалено %d осиротевших embeddings типа %s", len(orphan_ids), source_type)

    await session.commit()
    return stats
